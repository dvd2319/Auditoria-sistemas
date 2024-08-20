import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import matplotlib.pyplot as plt

# Definir las descripciones de las rúbricas específicas para cada pregunta
rubricas = {
    'Gobernanza y Liderazgo': {
        '¿Cómo demuestra la alta dirección su compromiso con la política antisoborno?': {
            1: 'La alta dirección no muestra un compromiso claro con la política antisoborno.',
            2: 'La alta dirección apoya la política antisoborno pero su participación es limitada.',
            3: 'La alta dirección muestra un compromiso activo con la política antisoborno.',
            4: 'La alta dirección lidera con el ejemplo y apoya activamente la política antisoborno.',
            5: 'La alta dirección está completamente comprometida y promueve activamente la política antisoborno en toda la organización.'
        },
        '¿Están claramente definidas las responsabilidades y autoridades para la gestión antisoborno?': {
            1: 'Las responsabilidades y autoridades no están definidas.',
            2: 'Las responsabilidades y autoridades están parcialmente definidas.',
            3: 'Las responsabilidades y autoridades están claramente definidas en la mayoría de los casos.',
            4: 'Las responsabilidades y autoridades están claramente definidas y comprendidas en toda la organización.',
            5: 'Las responsabilidades y autoridades están claramente definidas, comprendidas y se revisan regularmente para asegurar su eficacia.'
        }
    },
    'Gestión de Riesgos y Debida Diligencia': {
        '¿Se han identificado y documentado los riesgos de soborno en todas las áreas relevantes de la organización?': {
            1: 'No se han identificado ni documentado los riesgos de soborno.',
            2: 'Se han identificado y documentado algunos riesgos de soborno, pero de manera incompleta.',
            3: 'Se han identificado y documentado la mayoría de los riesgos de soborno.',
            4: 'Todos los riesgos de soborno relevantes están identificados y documentados.',
            5: 'Los riesgos de soborno están completamente identificados, documentados y se actualizan regularmente.'
        },
        '¿Se aplican consistentemente procedimientos de debida diligencia para evaluar a terceros y socios comerciales?': {
            1: 'No se aplican procedimientos de debida diligencia para terceros y socios comerciales.',
            2: 'Se aplican procedimientos de debida diligencia de manera limitada o inconsistente.',
            3: 'Se aplican procedimientos de debida diligencia en la mayoría de los casos.',
            4: 'Los procedimientos de debida diligencia se aplican consistentemente y cubren a todos los terceros y socios comerciales relevantes.',
            5: 'Los procedimientos de debida diligencia son exhaustivos, se aplican consistentemente y se revisan regularmente.'
        }
    },
    'Operaciones y Control de Procesos': {
        '¿Existen controles efectivos para la gestión de regalos, donaciones y hospitalidad que puedan influir en decisiones de negocio?': {
            1: 'No hay controles efectivos para la gestión de regalos, donaciones y hospitalidad.',
            2: 'Los controles para la gestión de regalos, donaciones y hospitalidad son limitados.',
            3: 'Existen controles efectivos para la mayoría de las situaciones que involucran regalos, donaciones y hospitalidad.',
            4: 'Los controles para la gestión de regalos, donaciones y hospitalidad son robustos y se aplican consistentemente.',
            5: 'Los controles son avanzados, con medidas proactivas adicionales para mitigar riesgos de soborno.'
        },
        '¿Se gestionan adecuadamente los riesgos de soborno en la cadena de suministro y contratación de proveedores?': {
            1: 'No se gestionan los riesgos de soborno en la cadena de suministro y contratación de proveedores.',
            2: 'Se gestionan de manera limitada o inconsistente los riesgos de soborno en la cadena de suministro.',
            3: 'La mayoría de los riesgos de soborno en la cadena de suministro y contratación de proveedores están gestionados.',
            4: 'Los riesgos de soborno en la cadena de suministro están adecuadamente gestionados y controlados.',
            5: 'La gestión de riesgos de soborno en la cadena de suministro es proactiva y se revisa regularmente.'
        }
    },
    'Capacitación y Concienciación': {
        '¿Se proporciona capacitación regular y efectiva sobre la política antisoborno a todos los empleados?': {
            1: 'No se proporciona capacitación sobre la política antisoborno.',
            2: 'La capacitación sobre la política antisoborno es limitada o insuficiente.',
            3: 'La mayoría de los empleados reciben capacitación regular sobre la política antisoborno.',
            4: 'Todos los empleados reciben capacitación regular y efectiva sobre la política antisoborno.',
            5: 'La capacitación es avanzada, personalizada y se actualiza regularmente para reflejar las mejores prácticas.'
        },
        '¿Existe una cultura organizacional que apoye activamente la política antisoborno?': {
            1: 'No existe una cultura organizacional que apoye la política antisoborno.',
            2: 'La cultura organizacional apoya la política antisoborno de manera limitada.',
            3: 'La cultura organizacional apoya activamente la política antisoborno en la mayoría de los casos.',
            4: 'La cultura organizacional apoya de manera consistente y activa la política antisoborno.',
            5: 'La cultura organizacional es completamente alineada con la política antisoborno y promueve la ética en todas las operaciones.'
        }
    },
    'Monitoreo, Revisión y Mejora Continua': {
        '¿Se realizan auditorías internas y monitoreo continuo del sistema de gestión antisoborno?': {
            1: 'No se realizan auditorías internas ni monitoreo del sistema de gestión antisoborno.',
            2: 'Se realizan auditorías y monitoreo de manera limitada o irregular.',
            3: 'Las auditorías internas y el monitoreo del sistema se realizan regularmente.',
            4: 'Las auditorías y el monitoreo son rigurosos y cumplen con todos los requisitos de la ISO 37001.',
            5: 'Las auditorías y el monitoreo son proactivos y superan los requisitos estándar, con mejoras continuas implementadas.'
        },
        '¿La alta dirección revisa periódicamente el sistema antisoborno para asegurar su efectividad y adecuación?': {
            1: 'La alta dirección no revisa el sistema antisoborno.',
            2: 'Las revisiones del sistema antisoborno por parte de la alta dirección son limitadas o irregulares.',
            3: 'La alta dirección revisa regularmente el sistema antisoborno.',
            4: 'Las revisiones son completas y aseguran la efectividad continua del sistema antisoborno.',
            5: 'Las revisiones son detalladas, frecuentes y conducen a mejoras continuas del sistema antisoborno.'
        }
    }
}

# Procesar las calificaciones y calcular los promedios
def procesar_calificaciones(calificaciones):
    promedios = {aspecto: sum(valores[1] for valores in lista) / len(lista) for aspecto, lista in calificaciones.items()}
    promedios_ponderados = {aspecto: (promedio / 5) * 20 for aspecto, promedio in promedios.items()}

    calificacion_final = sum(promedios_ponderados.values()) / len(promedios_ponderados) * 5
    return promedios_ponderados, calificacion_final

# Generar gráfico de barras utilizando matplotlib
def generar_grafico(promedios_ponderados):
    aspectos = list(promedios_ponderados.keys())
    valores = list(promedios_ponderados.values())

    plt.figure(figsize=(10, 6))
    plt.barh(aspectos, valores, color='skyblue')
    plt.xlabel('Nivel de Cumplimiento (sobre 20)')
    plt.title('Gráfico de Nivel de Cumplimiento por Aspecto')
    plt.xlim(0, 20)

    plt.tight_layout()
    plt.savefig('grafico_cumplimiento.png')
    st.pyplot(plt)

# Generar la conclusión general basada en la calificación final
def generar_conclusion(calificacion_final):
    if 0 <= calificacion_final <= 25:
        return ("El sistema de gestión antisoborno muestra una falta significativa de cumplimiento, "
                "con políticas y procedimientos insuficientes, controles deficientes, y una cultura organizacional "
                "que no apoya activamente la lucha contra el soborno. Es necesario implementar cambios profundos "
                "para alinearse con los requisitos de la ISO 37001.")
    elif 26 <= calificacion_final <= 50:
        return ("El sistema de gestión antisoborno tiene algunos controles y políticas en su lugar, pero estos no son suficientemente robustos "
                "o no se aplican consistentemente. Existen políticas y procedimientos documentados en algunas áreas, pero pueden estar "
                "desactualizados o no ser efectivos en la práctica. Los controles se implementan de manera limitada, y las revisiones se realizan "
                "de forma irregular.")
    elif 51 <= calificacion_final <= 75:
        return ("El sistema de gestión antisoborno ha implementado la mayoría de los controles requeridos por la norma ISO 37001. "
                "Las políticas y procedimientos están documentados y se revisan regularmente. La cultura organizacional apoya la lucha "
                "contra el soborno, y los controles son efectivos en la mayoría de los casos. No obstante, hay áreas que pueden mejorarse "
                "para alcanzar un nivel óptimo.")
    elif 76 <= calificacion_final <= 100:
        return ("El sistema de gestión antisoborno cumple completamente con los requisitos de la norma ISO 37001, y además implementa medidas adicionales "
                "que superan los estándares establecidos. Las políticas y procedimientos están completamente documentados y actualizados, y se revisan "
                "periódicamente. La cultura organizacional es fuerte y apoya activamente la lucha contra el soborno. Los controles son avanzados y se "
                "implementan de manera rigurosa.")
    else:
        return "Calificación no válida."

# Generar el informe en Word
def generar_informe_word(calificaciones, promedios_ponderados, calificacion_final, nombre_auditor, nombre_compania, fecha_evaluacion, destinatario, mensaje):
    document = Document()

    # Carátula
    document.add_heading('Informe de Evaluación de Cumplimiento de la Norma ISO 37001 (Sistema de Gestión Antisoborno)', 0)
    document.add_paragraph(f'Compañía Auditora: {nombre_compania}', style='Title')
    document.add_paragraph(f'Auditor: {nombre_auditor}', style='Heading 3')
    document.add_paragraph(f'Fecha de Evaluación: {fecha_evaluacion}', style='Heading 3')

    # Carta de introducción
    document.add_heading('Carta de Introducción', level=1)
    document.add_paragraph(f'Destinatario: {destinatario}', style='Heading 2')
    document.add_paragraph(mensaje)

    # Descripción del objetivo de la norma
    document.add_heading('Objetivo de la Norma ISO 37001', level=1)
    document.add_paragraph(
        "La norma ISO 37001 establece los requisitos para un sistema de gestión antisoborno, "
        "incluyendo las medidas que una organización debe implementar para prevenir, detectar y "
        "responder al soborno, así como cumplir con las leyes y regulaciones aplicables. "
        "El objetivo es crear una cultura de integridad, transparencia y cumplimiento dentro de la organización."
    )

    # Descripción de las dimensiones evaluadas
    document.add_heading('Dimensiones Evaluadas', level=1)
    document.add_paragraph(
        "A continuación se detallan las diferentes dimensiones evaluadas en este informe, junto con una breve descripción de cada una:"
    )

    dimensiones = {
        'Gobernanza y Liderazgo': "Evalúa el compromiso de la alta dirección y la claridad en la definición de responsabilidades y autoridades para la gestión antisoborno.",
        'Gestión de Riesgos y Debida Diligencia': "Evalúa la identificación, documentación y gestión de riesgos de soborno, así como la aplicación de debida diligencia para terceros y socios comerciales.",
        'Operaciones y Control de Procesos': "Evalúa la existencia y eficacia de controles para la gestión de regalos, donaciones, hospitalidad y la gestión de riesgos en la cadena de suministro.",
        'Capacitación y Concienciación': "Evalúa la efectividad de la capacitación sobre la política antisoborno y la cultura organizacional en apoyo a la lucha contra el soborno.",
        'Monitoreo, Revisión y Mejora Continua': "Evalúa la realización de auditorías internas, el monitoreo continuo y la revisión por parte de la alta dirección para asegurar la efectividad del sistema antisoborno."
    }

    for dimension, descripcion in dimensiones.items():
        document.add_heading(dimension, level=2)
        document.add_paragraph(descripcion)

    # Metodología de calificación
    document.add_heading('Metodología de Calificación', level=1)
    document.add_paragraph(
        "La evaluación se basa en una escala de 1 a 5, donde cada valor representa el nivel de cumplimiento de la norma:"
    )
    calificacion_metodologia = {
        1: "1 = No Cumple: No se realiza ninguna acción o la acción es insuficiente.",
        2: "2 = Cumple Parcialmente: Las acciones se realizan pero no con la frecuencia o efectividad requerida.",
        3: "3 = Cumple en Gran Medida: Las acciones se realizan regularmente y cumplen con la mayoría de los requisitos.",
        4: "4 = Cumple Totalmente: Las acciones cumplen con todos los requisitos establecidos.",
        5: "5 = Cumple y Supera las Expectativas: Se implementan medidas adicionales que superan los requisitos establecidos."
    }

    for key, value in calificacion_metodologia.items():
        document.add_paragraph(value)

    # Resultados de la evaluación
    document.add_heading('Resultados de la Evaluación', level=1)
    for aspecto, preguntas in calificaciones.items():
        document.add_heading(aspecto, level=2)
        for pregunta, calificacion in preguntas:
            descripcion = rubricas[aspecto][pregunta][calificacion]
            p = document.add_paragraph()
            p.add_run(f'{pregunta}: ').bold = True
            p.add_run(f'{calificacion} - {descripcion}')
        document.add_paragraph(f'Promedio del aspecto ({aspecto}): {promedios_ponderados[aspecto]:.2f} / 20')
        document.add_paragraph()

    document.add_paragraph(f'Calificación final del sistema antisoborno: {calificacion_final:.2f} / 100')
    document.add_paragraph()

    # Conclusión general
    conclusion = generar_conclusion(calificacion_final)
    document.add_heading('Conclusión General', level=1)
    document.add_paragraph(conclusion)
    document.add_paragraph()

    # Añadir gráfico de barras
    document.add_heading('Gráfico de Nivel de Cumplimiento por Aspecto', level=1)
    generar_grafico(promedios_ponderados)
    document.add_picture('grafico_cumplimiento.png', width=Inches(6))

    # Añadir pie de página
    section = document.sections[0]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = f'Compañía Auditora: {nombre_compania} - Fecha de Evaluación: {fecha_evaluacion}'
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.save('informe.docx')

    with open('informe.docx', 'rb') as f:
        st.download_button(label="Descargar Informe", data=f, file_name="informe.docx")

# Interfaz gráfica con streamlit
def main():
    st.title("Evaluación de Cumplimiento ISO 37001")
    st.write("Este aplicativo permite evaluar el cumplimiento de la norma ISO 37001 en la gestión antisoborno de su organización.")

    calificaciones = {key: [] for key in rubricas.keys()}

    for aspecto in rubricas.keys():
        st.subheader(aspecto)
        for pregunta in rubricas[aspecto].keys():
            calificacion = st.selectbox(pregunta, options=list(rubricas[aspecto][pregunta].keys()), format_func=lambda x: f'{x}: {rubricas[aspecto][pregunta][x]}')
            calificaciones[aspecto].append((pregunta, calificacion))

    if st.button("Generar Informe"):
        nombre_auditor = st.text_input("Nombre del Auditor")
        nombre_compania = st.text_input("Nombre de la Compañía Auditora")
        fecha_evaluacion = st.date_input("Fecha de Evaluación")
        destinatario = st.text_input("Destinatario del Informe")
        mensaje = st.text_area("Carta de Introducción")

        if not all([nombre_auditor, nombre_compania, fecha_evaluacion, destinatario, mensaje]):
            st.error("Debe completar todos los campos para generar el informe.")
        else:
            promedios_ponderados, calificacion_final = procesar_calificaciones(calificaciones)
            generar_informe_word(calificaciones, promedios_ponderados, calificacion_final,
                                 nombre_auditor, nombre_compania, fecha_evaluacion,
                                 destinatario, mensaje)

if __name__ == "__main__":
    main()
